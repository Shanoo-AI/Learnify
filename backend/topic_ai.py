import json
import os
import re

import docx
import requests


LLM_CONFIG = {
    "api_key": os.getenv("OPENROUTER_API_KEY"),
    "base_url": "https://openrouter.ai/api/v1/chat/completions",
    "model_name": os.getenv("OPENROUTER_MODEL", "deepseek/deepseek-chat"),
}

TOPIC_LIMIT = 80
AI_TIMEOUT_SECONDS = 12

_COURSE_CODE_RE = re.compile(r"\b[A-Z]{2,5}\s*[-_]?\s*\d{2,4}(?=\s|[A-Z]|\)|$)", re.IGNORECASE)
_WEEK_RE = re.compile(r"\b(week|unit|lecture|chapter|module|session)\s*#?\s*\d+\b", re.IGNORECASE)
_NUMBERED_RE = re.compile(r"^\s*(\d+|[ivxlcdm]+)[.)-]\s+", re.IGNORECASE)
_BULLET_RE = re.compile(r"^\s*[-*+\u2022\u25cf\u25aa\u25e6]\s*")
_WHITESPACE_RE = re.compile(r"\s+")
_QUESTION_MARKER_RE = re.compile(
    r"\b(?:question\s*(?:n[o0]\.?\s*)?#?\s*\d+|question\s*\d+|q\s*#?\s*\d+)\b",
    re.IGNORECASE,
)
_COMMAND_RE = re.compile(
    r"\b(apply|describe|discuss|differentiate|draw|design|explain|identify|outline|prepare|create|write|list|what|how|why)\b",
    re.IGNORECASE,
)
_ACADEMIC_TOPIC_RE = re.compile(
    r"\b("
    r"distance\s+vector\s+routing|selective\s+repeat\s+protocol|slotted\s+aloha|"
    r"attenuation|distortion|subnet(?:ting|s)?|subnet\s+mask|host\s+range|broadcast\s+id|"
    r"bit\s+rate|signal\s+level|router\s+(?:and|vs\.?|versus)\s+switch|"
    r"information\s+hiding|requirements?\s+traceability|non[-\s]+functional\s+requirements?|"
    r"project\s+scheduling|context\s+diagram|level-?\d+\s+dfd|dfd|state\s+machine\s+diagram|"
    r"class\s+diagram|technical\s+report|traffic\s+congestion|road\s+safety|"
    r"writing\s+procedures?|procedure\s+writing|tutorial\s+writing|step[-\s]+by[-\s]+step\s+tutorial"
    r")\b",
    re.IGNORECASE,
)

_NOISE_WORDS = {
    "assignment",
    "assignments",
    "assessment",
    "attendance",
    "bloom",
    "clo",
    "clos",
    "credit",
    "credits",
    "date",
    "exam",
    "final",
    "grading",
    "homework",
    "institute",
    "lab",
    "marks",
    "mid",
    "midterm",
    "office",
    "plo",
    "plos",
    "prerequisite",
    "quiz",
    "quizzes",
    "reference",
    "semester",
    "teacher",
    "teaching",
    "university",
}


def _clean_text(value):
    value = str(value or "").replace("\xa0", " ")
    value = value.replace("\r", "\n")
    value = re.sub(r"[;|]+", ", ", value)
    value = _BULLET_RE.sub("", value.strip())
    value = _NUMBERED_RE.sub("", value)
    value = _WHITESPACE_RE.sub(" ", value).strip(" ,.-:\t")
    return value


def _clean_subject(value):
    value = _clean_text(value)
    value = re.sub(r"\([^)]*\b(theory|lab|practical)\b[^)]*\)", "", value, flags=re.IGNORECASE)
    value = re.sub(r"\b\d+\s*\(\s*\d+\s*-\s*\d+\s*\)", "", value)
    value = _COURSE_CODE_RE.sub("", value)
    value = re.sub(r"\b(theory|practical|credit hours?|course outline|syllabus)\b", "", value, flags=re.IGNORECASE)
    value = re.sub(r"\(\s*[/,-]?\s*\)", "", value)
    value = value.strip(" /,-")
    return _clean_text(value)


def _topic_subject_key(value):
    value = _clean_subject(value)
    value = re.sub(r"\([^)]*\)", "", value)
    value = _COURSE_CODE_RE.sub("", value)
    return _clean_text(value).lower()


def _is_subject_noise(value):
    lowered = _clean_text(value).lower()
    if not lowered:
        return True
    if any(
        marker in lowered
        for marker in (
            "final term",
            "mid term",
            "examination",
            "registration",
            "total time",
            "maximum marks",
            "maximum points",
            "mor / eve",
            "mor/eve",
            "morning/evening",
        )
    ):
        return True
    if re.fullmatch(r"(?:bs|ms|bscs|bsse|bsit|bsai)[a-z/ -]*\d*(?:st|nd|rd|th)?", lowered):
        return True
    return False


def _dedupe(items):
    seen = set()
    result = []
    for item in items:
        cleaned = _clean_text(item)
        key = cleaned.lower()
        if cleaned and key not in seen:
            seen.add(key)
            result.append(cleaned)
    return result


def _unique_row_cells(row):
    return _dedupe(cell.text for cell in row.cells)


def _is_noise(text):
    text = _clean_text(text)
    if len(text) < 3:
        return True
    lowered = text.lower()
    words = re.findall(r"[a-z]+", lowered)
    if not words:
        return True
    if any(word in _NOISE_WORDS for word in words):
        # Keep real topics such as "software testing" while filtering admin rows.
        if not any(keeper in lowered for keeper in ("testing", "analysis", "design", "model", "network", "database")):
            return True
    if lowered in {"topics", "contents", "course contents", "week", "week #"}:
        return True
    if lowered in {"theory", "practical", "lecture", "lecture #", "domain", "bt level"}:
        return True
    if re.search(r"\b\d+\s*(%|marks?|hours?)\b", lowered):
        return True
    return False


def _topic_query(topic, subject_name):
    topic = _clean_text(topic)
    subject_name = _clean_subject(subject_name) or "Course Topic"
    subject_key = _topic_subject_key(subject_name)

    if not topic:
        return None
    if subject_key and subject_key in _topic_subject_key(topic):
        return topic
    return f"{topic} in {subject_name}"


def _split_topic_candidates(raw_text):
    raw_text = str(raw_text or "")
    raw_text = re.sub(r"[\u2022\u25cf\u25aa\u25e6]", "\n", raw_text)
    raw_text = re.sub(r"\b(?:and|with)\s+Number\s+of\s+Lectures.*", "", raw_text, flags=re.IGNORECASE)

    chunks = []
    for line in raw_text.splitlines():
        line = _clean_text(line)
        if not line:
            continue

        # Course outlines usually separate topics with commas or semicolons.
        parts = re.split(r",|;|\s{2,}", line)
        if len(parts) == 1 and len(line) > 130:
            parts = re.split(r"\.\s+", line)
        chunks.extend(parts)

    candidates = []
    for part in chunks:
        part = _clean_text(part)
        part = re.sub(r"^\b(week|unit|lecture|chapter|module|session)\s*#?\s*\d+\b[:\s-]*", "", part, flags=re.IGNORECASE)
        if _is_noise(part):
            continue
        if len(part) > 95:
            part = part[:95].rsplit(" ", 1)[0]
        candidates.append(part)

    return _dedupe(candidates)


def call_llm(system_prompt, user_prompt):
    headers = {
        "Content-Type": "application/json",
        "Authorization": f"Bearer {LLM_CONFIG['api_key']}",
    }
    payload = {
        "model": LLM_CONFIG["model_name"],
        "messages": [
            {"role": "system", "content": system_prompt},
            {"role": "user", "content": user_prompt},
        ],
        "temperature": 0.1,
    }

    try:
        response = requests.post(
            LLM_CONFIG["base_url"],
            headers=headers,
            json=payload,
            timeout=AI_TIMEOUT_SECONDS,
        )
        response.raise_for_status()
        return response.json()["choices"][0]["message"]["content"].strip()
    except Exception as e:
        print(f"AI topic helper unavailable: {e}")
        return None


def _subject_from_document(doc):
    header_text = []
    for table in doc.tables[:3]:
        for row in table.rows[:5]:
            header_text.extend(_unique_row_cells(row))
    header_text.extend(para.text for para in doc.paragraphs[:12])
    header_text = [
        line
        for raw in header_text
        for line in str(raw or "").splitlines()
        if _clean_text(line)
    ]

    for raw in header_text:
        text = _clean_text(raw)
        if not text:
            continue

        label_match = re.search(r"(course|subject)\s*(title|name)?\s*[:\-]\s*(.+)", text, re.IGNORECASE)
        if label_match:
            subject = _clean_subject(label_match.group(3))
            if subject:
                return subject

        if _COURSE_CODE_RE.search(text):
            subject = _clean_subject(text)
            if subject and not _is_subject_noise(subject):
                return subject

    for index, raw in enumerate(header_text):
        subject = _clean_subject(raw)
        if _COURSE_CODE_RE.fullmatch(_clean_text(raw)) and index > 0:
            previous_subject = _clean_subject(header_text[index - 1])
            if previous_subject and not _is_subject_noise(previous_subject):
                return previous_subject
        if subject and 2 <= len(subject.split()) <= 8 and not _is_noise(subject) and not _is_subject_noise(subject):
            return subject

    return "Course Topic"


def _subject_from_text(text):
    lines = [_clean_text(line) for line in str(text or "").splitlines()]
    lines = [line for line in lines[:60] if line]

    for line in lines:
        label_match = re.search(r"(course|subject)\s*(title|name)?\s*[:\-]\s*(.+)", line, re.IGNORECASE)
        if label_match:
            subject = _clean_subject(label_match.group(3))
            if subject:
                return subject

        if _COURSE_CODE_RE.search(line):
            subject = _clean_subject(line)
            if subject and len(subject) >= 4 and not _is_subject_noise(subject):
                return subject

    for index, line in enumerate(lines):
        if _COURSE_CODE_RE.fullmatch(line) and index > 0:
            subject = _clean_subject(lines[index - 1])
            if subject and not _is_noise(subject) and not _is_subject_noise(subject):
                return subject

    for line in lines:
        subject = _clean_subject(line)
        if subject and 2 <= len(subject.split()) <= 8 and not _is_noise(subject) and not _is_subject_noise(subject):
            return subject

    return "Course Topic"


def _document_text(doc):
    parts = []
    for para in doc.paragraphs:
        if para.text:
            parts.append(para.text)

    for table in doc.tables:
        for row in table.rows:
            cells = _unique_row_cells(row)
            if cells:
                parts.append(" | ".join(cells))

    return "\n".join(parts)


def detect_subject_via_ai(doc):
    subject = _subject_from_document(doc)
    if subject != "Course Topic":
        return subject

    raw_header_text = "\n".join(para.text for para in doc.paragraphs[:10])
    if len(doc.tables) > 0:
        raw_header_text += "\n" + "\n".join(
            " | ".join(_unique_row_cells(row)) for row in doc.tables[0].rows[:5]
        )

    sys_prompt = "You are a parser. Extract the specific Course/Subject Name from the text."
    usr_prompt = f"""
Analyze this raw header text from a course outline:
"{raw_header_text[:1000]}"

Return ONLY the clean subject name. Remove course codes, theory/lab labels, credits, and teacher info.
"""

    ai_subject = call_llm(sys_prompt, usr_prompt)
    return _clean_subject(ai_subject) if ai_subject else subject


def detect_subject_from_text_via_ai(text):
    subject = _subject_from_text(text)
    if subject != "Course Topic":
        return subject

    sys_prompt = "You are a parser. Extract the specific Course/Subject Name from the text."
    usr_prompt = f"""
Analyze this raw header text from a course outline:
"{str(text or '')[:1200]}"

Return ONLY the clean subject name. Remove course codes, theory/lab labels, credits, and teacher info.
"""

    ai_subject = call_llm(sys_prompt, usr_prompt)
    return _clean_subject(ai_subject) if ai_subject else subject


def _extract_json_list(text):
    if not text:
        return []

    clean_json = text.replace("```json", "").replace("```", "").strip()
    match = re.search(r"\[[\s\S]*\]", clean_json)
    if match:
        clean_json = match.group(0)

    try:
        parsed = json.loads(clean_json)
        if isinstance(parsed, list):
            return [str(item) for item in parsed if isinstance(item, str) and item.strip()]
    except Exception:
        pass

    return []


def extract_queries_via_ai(week, raw_text, subject_name):
    local_topics = _split_topic_candidates(raw_text)
    if local_topics:
        return [_topic_query(topic, subject_name) for topic in local_topics if _topic_query(topic, subject_name)]

    sys_prompt = "You are a YouTube Search Query Generator."
    usr_prompt = f"""
Subject: "{subject_name}"
Raw Text: "{raw_text}"

Extract learning sub-topics and output strictly a JSON list of YouTube search query strings.
Append "in {subject_name}" to each query. Ignore exams, dates, marks, CLOs, and admin text.
"""

    response = call_llm(sys_prompt, usr_prompt)
    queries = _extract_json_list(response)
    return queries or []


def _table_header_score(cells):
    lowered = " ".join(cells).lower()
    score = 0
    short_cells = [cell.lower() for cell in cells if len(cell) <= 70]
    short_header = " ".join(short_cells)

    if any(word in short_header for word in ("week", "unit", "lecture", "module", "session")):
        score += 2
    if any(word in short_header for word in ("topic", "content", "contents", "syllabus", "description", "detail")):
        score += 2
    if any(word in short_header for word in ("course", "covered")):
        score += 1
    if len(cells) == 1 and len(lowered) > 70:
        score = 0
    return score


def _topic_column_index(header_cells):
    for i, cell in enumerate(header_cells):
        lowered = cell.lower()
        if any(word in lowered for word in ("topic", "content", "contents", "syllabus", "description", "detail")):
            return i
    return 1 if len(header_cells) > 1 else 0


def _extract_from_tables(doc, subject_name):
    playlist = []

    for table in doc.tables:
        if len(table.rows) < 2:
            continue

        header_idx = None
        header_cells = []
        for idx, row in enumerate(table.rows[:3]):
            cells = _unique_row_cells(row)
            if _table_header_score(cells) >= 3:
                header_idx = idx
                header_cells = cells
                break

        if header_idx is None:
            first_body = _unique_row_cells(table.rows[1])
            first_row_looks_weekly = first_body and _WEEK_RE.search(first_body[0])
            if not first_row_looks_weekly:
                continue
            header_idx = 0
            header_cells = _unique_row_cells(table.rows[0])

        topic_idx = _topic_column_index(header_cells)

        for row in table.rows[header_idx + 1 :]:
            cells = _unique_row_cells(row)
            if not cells:
                continue

            if len(cells) > topic_idx:
                content_cells = [cells[topic_idx]]
            else:
                content_cells = cells[1:] if len(cells) > 1 else cells

            if len(cells) > 1 and _WEEK_RE.search(cells[0]):
                content_cells = cells[1:]

            content = " ".join(content_cells)
            if _is_noise(content):
                continue

            queries = extract_queries_via_ai(cells[0] if cells else "", content, subject_name)
            playlist.extend(queries)

        if playlist:
            break

    return _dedupe(playlist)


def _extract_from_paragraphs(doc, subject_name):
    lines = [_clean_text(para.text) for para in doc.paragraphs]
    lines = [line for line in lines if line and not _is_noise(line)]

    topic_lines = []
    capture = False
    for line in lines:
        lowered = line.lower()
        if any(marker in lowered for marker in ("topics covered", "course contents", "weekly plan", "syllabus", "lecture plan")):
            capture = True
            continue

        if capture and any(marker in lowered for marker in ("course learning outcomes", "assessment", "grading", "reference")):
            break

        if capture or _WEEK_RE.search(line) or _NUMBERED_RE.search(line):
            topic_lines.append(line)

    if not topic_lines:
        topic_lines = lines[:30]

    topics = []
    for line in topic_lines:
        topics.extend(_split_topic_candidates(line))

    return _dedupe(_topic_query(topic, subject_name) for topic in topics if topic)[:TOPIC_LIMIT]


def _extract_text_from_pdf(file_path):
    try:
        import fitz

        pages = []
        with fitz.open(file_path) as pdf:
            for page in pdf:
                page_text = (page.get_text("text") or "").strip()
                if page_text:
                    pages.append(page_text)
        if pages:
            return "\n".join(pages)
    except Exception as exc:
        print(f"PyMuPDF PDF extraction failed: {exc}")

    try:
        import pdfplumber

        pages = []
        with pdfplumber.open(file_path) as pdf:
            for page in pdf.pages:
                page_text = (page.extract_text() or "").strip()
                if page_text:
                    pages.append(page_text)
        return "\n".join(pages)
    except Exception as exc:
        print(f"pdfplumber PDF extraction failed: {exc}")
        return ""


def _extract_text_from_legacy_doc(file_path):
    """Best-effort extraction for old binary .doc files without Word installed."""
    try:
        with open(file_path, "rb") as handle:
            data = handle.read()
    except Exception as exc:
        print(f"Legacy DOC read failed: {exc}")
        return ""

    candidates = []
    for encoding in ("utf-16le", "latin1"):
        decoded = data.decode(encoding, errors="ignore")
        printable = "".join(
            char if char in "\r\n\t" or 32 <= ord(char) <= 126 else " "
            for char in decoded
        )
        printable = printable.replace("\r", "\n")
        printable = re.sub(r"[ \t]{2,}", " ", printable)
        printable = re.sub(r"\n\s+", "\n", printable).strip()
        score = len(re.findall(r"\b(question|course|subject|marks|semester|diagram|explain)\b", printable, re.IGNORECASE))
        score += len(re.findall(r"[A-Za-z]{4,}", printable)) // 25
        candidates.append((score, printable))

    candidates.sort(key=lambda item: item[0], reverse=True)
    return candidates[0][1] if candidates else ""


def _looks_like_exam_text(text):
    text = str(text or "")
    question_count = len(_QUESTION_MARKER_RE.findall(text))
    lowered = text.lower()
    return question_count >= 2 or (
        question_count >= 1 and any(word in lowered for word in ("marks", "clo", "examination", "paper"))
    )


def _clean_exam_question(text):
    text = re.sub(_QUESTION_MARKER_RE, " ", str(text or ""))
    text = re.sub(r"\([^)]*\b(?:marks?|clo|points?)\b[^)]*\)", " ", text, flags=re.IGNORECASE)
    text = re.sub(r"\[[^\]]*\b(?:marks?|clo|points?)\b[^\]]*\]", " ", text, flags=re.IGNORECASE)
    text = re.sub(r"\b(?:marks?|clo|points?)\s*[-:=]?\s*\d+(?:\+\d+)*\b", " ", text, flags=re.IGNORECASE)
    text = re.sub(r"\b[a-d]\)\s*", ". ", text, flags=re.IGNORECASE)
    return _clean_text(text)


def _trim_exam_topic(text):
    text = _clean_text(text)
    text = re.sub(
        r"^(?:based on this scenario|based on the principles of|answer the following|following|the below questions|as a [^,]+,\s*)",
        "",
        text,
        flags=re.IGNORECASE,
    )
    text = re.sub(
        r"\b(?:and answer the following|answer the below questions|with the help of diagram|support your answer.*|at least.*|using the essential.*|following all key.*|initial state.*|fully convergence.*|identifying key entities.*|ensuring that you include.*)\b.*",
        "",
        text,
        flags=re.IGNORECASE,
    )
    text = re.sub(r"^(?:you are required to|you have been assigned to|you have been hired as|your task is to)\s+", "", text, flags=re.IGNORECASE)
    text = re.sub(r"^(?:the|a|an)\s+", "", text, flags=re.IGNORECASE)
    text = text.strip(" ,.-:?")
    if len(text) > 120:
        text = text[:120].rsplit(" ", 1)[0]
    return _clean_text(text)


def _extract_diagram_topics(question):
    topics = []
    for match in re.finditer(
        r"\b(?:level-?\d+\s*)?(?:context\s+diagram|dfd|state\s+machine\s+diagram|class\s+diagram|use\s+case\s+diagram|sequence\s+diagram|activity\s+diagram|er\s+diagram)\b",
        question,
        re.IGNORECASE,
    ):
        topics.append(_clean_text(match.group(0)))
    return topics


def _extract_academic_topics(question):
    topics = []
    for match in _ACADEMIC_TOPIC_RE.finditer(question):
        topic = _clean_text(match.group(1))
        topic = re.sub(r"\s+(?:and|vs\.?|versus)\s+", " vs ", topic, flags=re.IGNORECASE)
        if topic.lower() == "subnets":
            topic = "subnetting"
        topics.append(topic)
    return topics


def _extract_exam_topics_from_text(text, subject_name):
    text = str(text or "").replace("\r", "\n")
    markers = list(_QUESTION_MARKER_RE.finditer(text))
    if not markers:
        return []

    chunks = []
    for index, marker in enumerate(markers):
        start = marker.start()
        end = markers[index + 1].start() if index + 1 < len(markers) else len(text)
        chunks.append(text[start:end])

    topics = []
    for chunk in chunks:
        question = _clean_exam_question(chunk)
        if _is_noise(question):
            continue

        direct_topics = _extract_diagram_topics(question)
        direct_topics.extend(_extract_academic_topics(question))

        if direct_topics:
            topics.extend(direct_topics)
            continue

        clauses = re.split(r"(?:\.\s+|\?\s+|\n+|;\s+)", question)
        useful_clauses = [clause for clause in clauses if _COMMAND_RE.search(clause)]
        if not useful_clauses:
            useful_clauses = clauses[:2]

        for clause in useful_clauses[:4]:
            clause = _clean_text(clause)
            if _is_noise(clause):
                continue

            command_match = _COMMAND_RE.search(clause)
            if command_match:
                clause = clause[command_match.end():]

            phrase_match = re.search(
                r"\b(?:on|about|between|for|of)\s+([A-Z][A-Za-z0-9 /&-]{3,80})",
                clause,
            )
            if phrase_match:
                clause = phrase_match.group(1)

            clause = re.sub(
                r"^(?:the|a|an|on|about|between|how|why|what|to|of|main|key|three)\s+",
                "",
                clause,
                flags=re.IGNORECASE,
            )
            topic = _trim_exam_topic(clause)
            if topic and not _is_noise(topic):
                topics.append(topic)

    return _dedupe(_topic_query(topic, subject_name) for topic in topics if topic)[:TOPIC_LIMIT]


def _extract_from_plain_text(text, subject_name):
    lines = [_clean_text(line) for line in str(text or "").splitlines()]
    lines = [line for line in lines if line]

    topic_lines = []
    capture = False

    for line in lines:
        lowered = line.lower()

        if "week" in lowered and any(
            marker in lowered for marker in ("theory", "topic", "content", "lecture")
        ):
            capture = True
            continue

        if any(marker in lowered for marker in ("topics covered", "course contents", "weekly plan", "syllabus")):
            capture = True
            continue

        if capture and any(marker in lowered for marker in ("course learning outcomes", "assessment", "grading", "reference materials")):
            continue

        if capture:
            if re.fullmatch(r"\d+\s*&\s*\d+", line):
                continue
            if re.fullmatch(r"\d+\s*(?:-\s*\d+)?", line):
                continue
            if _WEEK_RE.fullmatch(line):
                continue
            if _is_noise(line):
                continue
            topic_lines.append(line)

    if not topic_lines:
        for line in lines:
            if _WEEK_RE.search(line) or _NUMBERED_RE.search(line):
                topic_lines.append(line)

    topics = []
    for line in topic_lines:
        topics.extend(_split_topic_candidates(line))

    return _dedupe(_topic_query(topic, subject_name) for topic in topics if topic)[:TOPIC_LIMIT]


def process_smartly(file_path):
    extension = os.path.splitext(str(file_path))[1].lower()

    print("Analyzing course outline structure...")

    if extension == ".pdf":
        text = _extract_text_from_pdf(file_path)
        if not text.strip():
            print("No readable text found in PDF.")
            return []

        subject_name = detect_subject_from_text_via_ai(text)
        print(f"Detected subject: {subject_name}")
        if _looks_like_exam_text(text):
            exam_topics = _extract_exam_topics_from_text(text, subject_name)
            if exam_topics:
                return exam_topics
        return _extract_from_plain_text(text, subject_name)

    if extension == ".doc":
        text = _extract_text_from_legacy_doc(file_path)
        if not text.strip():
            print("No readable text found in DOC.")
            return []

        subject_name = detect_subject_from_text_via_ai(text)
        print(f"Detected subject: {subject_name}")
        if _looks_like_exam_text(text):
            exam_topics = _extract_exam_topics_from_text(text, subject_name)
            if exam_topics:
                return exam_topics
        return _extract_from_plain_text(text, subject_name)

    doc = docx.Document(file_path)
    text = _document_text(doc)

    doc_subject = detect_subject_via_ai(doc)
    text_subject = detect_subject_from_text_via_ai(text)
    subject_name = text_subject if _looks_like_exam_text(text) and _is_subject_noise(doc_subject) else doc_subject
    print(f"Detected subject: {subject_name}")

    if _looks_like_exam_text(text):
        exam_topics = _extract_exam_topics_from_text(text, subject_name)
        if exam_topics:
            return exam_topics

    playlist = _extract_from_tables(doc, subject_name)
    if not playlist:
        playlist = _extract_from_paragraphs(doc, subject_name)

    return playlist[:TOPIC_LIMIT]


if __name__ == "__main__":
    final_qs = process_smartly("SE_outline.docx")
    print("\nSAMPLE OUTPUT:")
    print(final_qs[:5])
