import re
import json
import uuid
import random
import os
import time
from pathlib import Path
from datetime import datetime
import chromadb
import requests
from pymongo import MongoClient
from flask import Flask, request, jsonify, session
from flask_cors import CORS
from itsdangerous import BadSignature, SignatureExpired, URLSafeTimedSerializer
from werkzeug.utils import secure_filename

# ========== FLASK SETUP ==========
app = Flask(__name__)
app.secret_key = os.getenv("SECRET_KEY", "dev-secret-key-change-me")
auth_serializer = URLSafeTimedSerializer(app.secret_key, salt="learnify-auth")
_cors_origins = [
    origin.strip()
    for origin in os.getenv("CORS_ORIGINS", os.getenv("FRONTEND_URL", "")).split(",")
    if origin.strip()
]
CORS(
    app,
    supports_credentials=True,
    origins=_cors_origins or [r"http://localhost(:\d+)?", r"http://127\.0\.0\.1(:\d+)?"],
)


def load_auth_token(token):
    try:
        data = auth_serializer.loads(token, max_age=60 * 60 * 24 * 30)
        return data.get("user")
    except (BadSignature, SignatureExpired):
        return None


@app.before_request
def load_bearer_session():
    if session.get("logged_in"):
        return

    auth_header = request.headers.get("Authorization", "")
    if not auth_header.startswith("Bearer "):
        return

    user = load_auth_token(auth_header.removeprefix("Bearer ").strip())
    if user:
        session["user"] = user
        session["logged_in"] = True

# ========== DATABASE SETUP ==========
# MongoDB
mongo_client = MongoClient(os.getenv("MONGO_URI", "mongodb://localhost:27017/"))
db = mongo_client["learnify"]
chats_collection = db["chats"]
sessions_collection = db["sessions"]

# ChromaDB - Persistent storage
chroma_client = chromadb.PersistentClient(path=os.getenv("CHROMA_DB_PATH", "./chroma_db"))
collection = chroma_client.get_or_create_collection("course_topics")
session_docs_collection = chroma_client.get_or_create_collection("session_documents")

UPLOADS_FOLDER = Path(__file__).resolve().parent / "uploads"
UPLOADS_FOLDER.mkdir(exist_ok=True)
ALLOWED_DOCUMENT_EXTENSIONS = {".pdf", ".docx", ".txt"}

# ========== GEMINI CONFIG ==========
# Paste your Google Generative AI key here later, or set GEMINI_API_KEY
# in your environment to avoid keeping secrets in code.
GEMINI_API_KEY = os.getenv("GEMINI_API_KEY")
GEMINI_MODELS = [
    "gemini-2.5-flash",
    "gemini-2.5-flash-lite",
]

# ========== EMERGENCY GROQ FALLBACK (COMMENTED) ==========
# If Gemini causes issues during a presentation, uncomment the import/config/call
# below and comment the Gemini call instead.
#
# from groq import Groq
# GROQ_API_KEY = "YOUR_GROQ_API_KEY_HERE"
# groq_client = Groq(api_key=GROQ_API_KEY)
# GROQ_MODEL = "llama-3.1-8b-instant"
#
# def generate_with_groq(system_prompt, user_message):
#     chat_completion = groq_client.chat.completions.create(
#         messages=[
#             {"role": "system", "content": system_prompt},
#             {"role": "user", "content": user_message}
#         ],
#         model=GROQ_MODEL,
#         temperature=0.7,
#         max_tokens=2000
#     )
#     return chat_completion.choices[0].message.content


def build_gemini_api_url(model_name):
    return (
        f"https://generativelanguage.googleapis.com/v1beta/models/"
        f"{model_name}:generateContent"
    )


def generate_with_gemini(system_prompt, user_message):
    """Send a text-only request to the Gemini API."""
    if not GEMINI_API_KEY or GEMINI_API_KEY == "YOUR_GEMINI_API_KEY_HERE":
        raise ValueError(
            "Gemini API key not configured. Set GEMINI_API_KEY or replace "
            "YOUR_GEMINI_API_KEY_HERE in backend/mentorbot/backend/app.py."
        )

    last_error = None

    for model_name in GEMINI_MODELS:
        payload = {
            "system_instruction": {
                "parts": [{"text": system_prompt}]
            },
            "contents": [
                {
                    "role": "user",
                    "parts": [{"text": user_message}]
                }
            ],
            "generationConfig": {
                "temperature": 0.7,
                "maxOutputTokens": 2000,
                "responseMimeType": "text/plain",
                # Gemini 2.5 Flash thinks by default; forcing zero reduces delay/load.
                "thinkingConfig": {
                    "thinkingBudget": 0
                }
            }
        }

        for attempt in range(2):
            try:
                response = requests.post(
                    build_gemini_api_url(model_name),
                    headers={
                        "x-goog-api-key": GEMINI_API_KEY,
                        "Content-Type": "application/json",
                    },
                    json=payload,
                    timeout=45,
                )

                if not response.ok:
                    try:
                        error_data = response.json()
                        error_message = error_data.get("error", {}).get("message", response.text)
                    except ValueError:
                        error_message = response.text

                    transient_status = response.status_code in (429, 500, 503)
                    high_demand = "high demand" in error_message.lower() or "resource_exhausted" in error_message.lower()

                    if attempt == 0 and (transient_status or high_demand):
                        time.sleep(1.2)
                        continue

                    raise RuntimeError(
                        f"Gemini API request failed on {model_name}: {error_message}"
                    )

                data = response.json()
                candidates = data.get("candidates", [])
                if not candidates:
                    raise RuntimeError(f"Gemini API returned no candidates on {model_name}.")

                parts = candidates[0].get("content", {}).get("parts", [])
                text_parts = [part.get("text", "") for part in parts if part.get("text")]
                assistant_reply = "\n".join(text_parts).strip()

                if assistant_reply:
                    return assistant_reply

                raise RuntimeError(f"Gemini API returned an empty response on {model_name}.")
            except (requests.RequestException, RuntimeError) as exc:
                last_error = exc
                if attempt == 0:
                    time.sleep(0.8)
                    continue
                break

    raise RuntimeError(str(last_error) if last_error else "Gemini API request failed.")

# ========== COURSE EXTRACTION ==========
def extract_course_data(text):
    """Extract course information from course outline text"""
    data = {
        "course_name": "",
        "credit_hours": "",
        "prerequisites": "",
        "weekly_topics": []
    }
    
    course_match = re.search(r'([A-Z]{2,4}-\d{3})\s+(.+?)(?:\n|\*\*|Credit)', text, re.IGNORECASE)
    if course_match:
        data["course_name"] = f"{course_match.group(1)} {course_match.group(2)}".strip()
    
    credit_match = re.search(r'(?:Credit\s+Hours?)[:\s*]+(\d+(?:\(\d+-\d+\))?)', text, re.IGNORECASE)
    if credit_match:
        data["credit_hours"] = credit_match.group(1)
    else:
        credit_match2 = re.search(r'\*\*(\d+\(\d+-\d+\))\*\*', text)
        if credit_match2:
            data["credit_hours"] = credit_match2.group(1)
    
    prereq_match = re.search(r'Prerequisites?[:\s*]+(.+?)(?:\n|Teacher|Office|\*\*)', text, re.IGNORECASE)
    if prereq_match:
        data["prerequisites"] = prereq_match.group(1).strip()
    
    week_pattern = r'Week\s+(\d+)\s+(.+?)(?=Week\s+\d+|\*\*Mid|\*\*Final|$)'
    weeks = re.findall(week_pattern, text, re.IGNORECASE | re.DOTALL)
    
    for week_num, content in weeks:
        topics = re.sub(r'\s+', ' ', content).strip()
        topics = re.sub(r'\*\*', '', topics)
        topics = topics.split('Lect-I')[-1] if 'Lect-I' in topics else topics
        topics = topics.split('L ect-I')[-1] if 'L ect-I' in topics else topics
        topics = topics.split('Le ct-II')[0] if 'Le ct-II' in topics else topics
        topics = topics[:300] if len(topics) > 300 else topics
        
        data["weekly_topics"].append({
            "week": int(week_num),
            "topics": topics.strip()
        })
    
    return data

def read_docx(file_path):
    """Read text from a .docx file"""
    try:
        from docx import Document
        doc = Document(file_path)
        full_text = []
        for para in doc.paragraphs:
            full_text.append(para.text)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    full_text.append(cell.text)
        return '\n'.join(full_text)
    except ImportError:
        print("\nError: python-docx library not installed!")
        print("Please install it using: pip install python-docx\n")
        exit(1)


def extract_text_from_pdf(file_path):
    """Extract text from PDF using PyMuPDF first, pdfplumber second."""
    try:
        import fitz

        extracted = []
        with fitz.open(file_path) as doc:
            for page in doc:
                page_text = (page.get_text("text") or "").strip()
                if page_text:
                    extracted.append(page_text)
        text = "\n".join(extracted).strip()
        if text:
            return text
    except ImportError:
        pass
    except Exception:
        pass

    try:
        import pdfplumber

        extracted = []
        with pdfplumber.open(file_path) as pdf:
            for page in pdf.pages:
                page_text = (page.extract_text() or "").strip()
                if page_text:
                    extracted.append(page_text)
        text = "\n".join(extracted).strip()
        if text:
            return text
    except ImportError:
        raise RuntimeError(
            "PDF support requires PyMuPDF or pdfplumber. Install one of them."
        )

    raise RuntimeError("Could not extract readable text from PDF.")


def extract_document_text(file_path):
    """Extract text from supported document types."""
    suffix = Path(file_path).suffix.lower()
    if suffix == ".txt":
        return Path(file_path).read_text(encoding="utf-8", errors="ignore").strip()
    if suffix == ".docx":
        return read_docx(file_path).strip()
    if suffix == ".pdf":
        return extract_text_from_pdf(file_path).strip()
    raise RuntimeError("Unsupported file type.")


def chunk_text(text, chunk_size=1200, overlap=150):
    """Chunk large text into overlapping slices for retrieval."""
    normalized = re.sub(r"\s+", " ", text).strip()
    if not normalized:
        return []

    chunks = []
    start = 0
    text_length = len(normalized)

    while start < text_length:
        end = min(start + chunk_size, text_length)
        chunk = normalized[start:end].strip()
        if chunk:
            chunks.append(chunk)
        if end >= text_length:
            break
        start = max(end - overlap, start + 1)

    return chunks


def get_session_documents(session_id, user_id):
    chat_session = sessions_collection.find_one(
        session_owner_filter(session_id, user_id),
        {"documents": 1, "_id": 0},
    )
    return chat_session.get("documents", []) if chat_session else []


def add_document_to_session(session_id, user_id, document_info):
    sessions_collection.update_one(
        session_owner_filter(session_id, user_id),
        {
            "$push": {"documents": document_info},
            "$set": {"updated_at": datetime.now()},
        },
        upsert=False,
    )


def remove_session_document_embeddings(session_id, document_id=None):
    try:
        results = session_docs_collection.get(where={"session_id": session_id})
        ids = []
        metadatas = results.get("metadatas") or []
        result_ids = results.get("ids") or []
        for result_id, metadata in zip(result_ids, metadatas):
            if not document_id or metadata.get("document_id") == document_id:
                ids.append(result_id)
        if ids:
            session_docs_collection.delete(ids=ids)
    except Exception as exc:
        print(f"Session document cleanup warning: {exc}")


def build_document_context(session_id, user_message, max_chunks=4):
    """Retrieve relevant uploaded document chunks for the current session."""
    try:
        results = session_docs_collection.query(
            query_texts=[user_message],
            n_results=max_chunks,
            where={"session_id": session_id},
        )
    except Exception as exc:
        print(f"Session document query warning: {exc}")
        return ""

    documents = results.get("documents") or []
    metadatas = results.get("metadatas") or []

    if not documents or not documents[0]:
        return ""

    context_lines = ["**Relevant uploaded document excerpts:**"]
    for index, chunk in enumerate(documents[0]):
        metadata = metadatas[0][index] if metadatas and metadatas[0] else {}
        file_name = metadata.get("file_name", "Uploaded document")
        chunk_no = metadata.get("chunk_index", index + 1)
        context_lines.append(f"\n{file_name} - Chunk {chunk_no}:\n{chunk}")

    return "\n".join(context_lines)

# ========== INTENT DETECTION ==========
def detect_intent(message):
    """Classify user intent while avoiding accidental substring matches."""
    msg_lower = message.lower().strip()
    word_count = len(msg_lower.split())

    def contains_phrase(phrases):
        for phrase in phrases:
            pattern = r'(?<!\w)' + re.escape(phrase) + r'(?!\w)'
            if re.search(pattern, msg_lower):
                return True
        return False
    
    # Greetings (no RAG needed)
    greetings = [
        'hi', 'hello', 'salam', 'assalam', 'hey', 'sup', 'greetings',
        'howdy', 'good morning', 'good afternoon', 'good evening',
        'kesa hai', 'kya hal', 'how are you', 'kese ho', 'kaise ho'
    ]
    if contains_phrase(greetings) and word_count <= 4:
        return "greeting"
    
    # Farewells
    farewells = [
        'bye', 'goodbye', 'khuda hafiz', 'allah hafiz', 'see you',
        'good night', 'goodnight', 'take care', 'farewell'
    ]
    if contains_phrase(farewells) and word_count <= 5:
        return "farewell"
    
    # Acknowledgments
    acks = [
        'thanks', 'thank you', 'shukriya', 'ok', 'okay', 'acha', 'theek hai',
        'got it', 'understood', 'samajh gaya', 'samajh gayi', 'awesome',
        'great', 'perfect', 'makes sense'
    ]
    if contains_phrase(acks) and word_count <= 5:
        return "acknowledgment"
    
    # Academic question indicators
    question_words = [
        'kya', 'kese', 'kaise', 'explain', 'bta', 'batao', 'batayen',
        'what', 'how', 'why', 'when', 'where', 'who', 'define', 'elaborate',
        'help', 'samjhao', 'topics', 'topic', 'code', 'algorithm', 'function',
        'error', 'bug', 'course', 'week', 'lecture'
    ]
    if contains_phrase(question_words) or '?' in msg_lower:
        return "question"
    
    # Default: treat as question if unsure (safer for academic bot)
    return "question"

def handle_casual_conversation(message, intent):
    """Handle greetings and casual talk without RAG"""
    
    responses = {
        "greeting": [
            "Hello! Aaj kis topic mein help chahiye?",
            "Hi there! Batao, studies mein kya discuss karna hai?",
            "Assalam o Alaikum! Main ready hoon, kis concept par baat karein?",
            "Hey! Koi question ya topic hai jo explain karna hai?"
        ],
        "farewell": [
            "Allah hafiz! Jab bhi help chahiye ho wapas aa jana.",
            "Goodbye! Happy learning, aur practice zaroor karna.",
            "Take care! Koi aur question ho to pooch lena.",
            "Bye for now! Main yahin hoon jab next doubt aaye."
        ],
        "acknowledgment": [
            "You're welcome! Aur kuch clarify karna ho to batao.",
            "Glad I could help! Next topic kya dekhna hai?",
            "Perfect! Koi aur concept stuck ho to pooch lena.",
            "Sounds good! Need more examples or practice questions?"
        ]
    }
    
    return random.choice(responses.get(intent, ["Understood. How can I help further?"]))

def detect_subject_from_text(text):
    """Intelligently detect which subject user is talking about"""
    text_lower = text.lower()
    
    # Algorithm keywords
    algo_keywords = ['algo', 'algorithm', 'sorting', 'asymptotic', 'recursion', 'dynamic programming', 
                     'greedy', 'graph', 'heap', 'complexity', 'big-o', 'merge sort', 'quick sort',
                     'csc-201', 'csc 201', 'data structure']
    
    # Marketing keywords  
    marketing_keywords = ['marketing', 'customer', 'brand', 'product', 'promotion', 'price', 
                          'distribution', 'mgt-351', 'mgt 351', 'sales', 'advertising']
    
    algo_score = sum(1 for keyword in algo_keywords if keyword in text_lower)
    marketing_score = sum(1 for keyword in marketing_keywords if keyword in text_lower)
    
    if algo_score > marketing_score:
        return "CSC-201"
    elif marketing_score > algo_score:
        return "MGT-351"
    return None

def process_course_files():
    """Extract data from all course files and store in ChromaDB"""
    files_folder = Path("files")
    
    if not files_folder.exists():
        print(f"Error: 'files' folder not found!")
        return None
    
    files = list(files_folder.glob("*.txt")) + list(files_folder.glob("*.docx"))
    
    if not files:
        print("No .txt or .docx files found in 'files' folder!")
        return None
    
    print(f"Found {len(files)} file(s) in 'files' folder\n")
    print("="*80)
    
    all_courses = []
    
    for file_path in files:
        print(f"\nProcessing: {file_path.name}")
        print("-"*80)
        
        try:
            if file_path.suffix.lower() == '.docx':
                course_text = read_docx(file_path)
            else:
                with open(file_path, 'r', encoding='utf-8') as f:
                    course_text = f.read()
            
            result = extract_course_data(course_text)
            all_courses.append(result)
            
            # Extract course code for filtering
            course_code = result["course_name"].split()[0] if result["course_name"] else "UNKNOWN"
            
            # Store in ChromaDB with course code metadata
            for week in result["weekly_topics"]:
                doc_id = f"{course_code}_week_{week['week']}"
                
                try:
                    existing = collection.get(ids=[doc_id])
                    if not existing['ids']:
                        collection.add(
                            documents=[week["topics"]],
                            metadatas=[{
                                "course": result["course_name"], 
                                "course_code": course_code,
                                "week": week["week"]
                            }],
                            ids=[doc_id]
                        )
                except:
                    collection.add(
                        documents=[week["topics"]],
                        metadatas=[{
                            "course": result["course_name"],
                            "course_code": course_code, 
                            "week": week["week"]
                        }],
                        ids=[doc_id]
                    )
            
            print(json.dumps(result, indent=2, ensure_ascii=False))
            print("\n" + "="*80)
            
        except Exception as e:
            print(f"Error processing {file_path.name}: {e}")
            print("="*80)
    
    # Save courses info
    with open('courses.json', 'w', encoding='utf-8') as f:
        json.dump(all_courses, f, indent=2, ensure_ascii=False)
    
    print(f"\n✓ Total courses loaded: {len(all_courses)}")
    print(f"✓ ChromaDB indexed: {collection.count()} topics")
    
    return all_courses

# ========== CONVERSATION MANAGEMENT ==========
def get_current_user_id():
    """Return the logged-in app user for MentorBot ownership checks."""
    if not session.get("logged_in") or not session.get("user"):
        return None
    return str(session.get("user"))


def require_current_user():
    user_id = get_current_user_id()
    if not user_id:
        return None, (jsonify({"error": "Please login first"}), 401)
    return user_id, None


def session_owner_filter(session_id, user_id):
    return {"session_id": session_id, "user_id": user_id}


def get_or_create_session(session_id, user_id):
    """Get existing session or create new one"""
    existing = sessions_collection.find_one({"session_id": session_id})
    if existing and existing.get("user_id") != user_id:
        return None
    
    if not existing:
        existing = {
            "session_id": session_id,
            "user_id": user_id,
            "created_at": datetime.now(),
            "messages": [],
            "documents": [],
            "context": {
                "current_subject": None,
                "topics_discussed": []
            }
        }
        sessions_collection.insert_one(existing)
    
    return existing

def update_session_context(session_id, user_id, user_message, assistant_reply, detected_subject):
    """Update session with new message and context"""
    sessions_collection.update_one(
        session_owner_filter(session_id, user_id),
        {
            "$push": {
                "messages": {
                    "$each": [
                        {"role": "user", "content": user_message, "timestamp": datetime.now()},
                        {"role": "assistant", "content": assistant_reply, "timestamp": datetime.now()}
                    ]
                }
            },
            "$set": {
                "context.current_subject": detected_subject,
                "updated_at": datetime.now()
            }
        }
    )

def get_conversation_history(session_id, user_id, last_n=5):
    """Get last N messages from conversation"""
    chat_session = sessions_collection.find_one(session_owner_filter(session_id, user_id))
    if chat_session and "messages" in chat_session:
        return chat_session["messages"][-last_n:]
    return []

# ========== INTELLIGENT CHAT ==========
SYSTEM_PROMPT = """You are Learnify, an intelligent AI tutor for university students in Pakistan. 
**VERY STRICT MUST FOLLOW INSTRUCTION:**
- Always start conversation in English
**LANGUAGE RULES (CRITICAL):**
- Always start conversation in English unless the student uses Roman Urdu, then reply in Roman Urdu.
- reply in the language that the student uses in their question (if they use Roman Urdu, reply in Roman Urdu; if they use English, reply in English)
- use Roman Urdu (Urdu written in English) mixed with English
- NEVER use Hindi words or Devanagari script
- While talking in Roman Urdu Use Pakistani Urdu expressions: "jani", "yaar", "bhai", "acha", "theek hai"
- Examples: "Bilkul theek hai jani", "Samajh gaya yaar", "Acha to ye topic..."

**PERSONALITY:**
- Friendly, patient, and encouraging like a senior student helping juniors
- Use casual tone: "dekho jani", "simple hai yaar", "tension na lo"
- Give step-by-step explanations
- Use examples from Pakistani context when relevant

**RESPONSE RULES:**
1. **For greetings/casual talk:** Keep it short, friendly, and ask what they need help with
2. **For academic questions:** Use clear headings, bullet points, and detailed explanations
3. **For topic explanations:** Explain thoroughly with examples from course content
4. **For vague questions:** Ask clarifying questions to understand what they need
5. **Formatting:** Always format academic answers in clean Markdown-style text

**OUTPUT FORMAT FOR ACADEMIC ANSWERS:**
- Start with a short heading using `##`
- Use short paragraphs, not one huge block
- Use `-` bullet points for key points
- Use `1.` numbered steps when explaining a process
- Use `**bold**` for important terms
- If relevant, end with a short `## Quick Summary`
- Keep spacing clean between sections

**CRITICAL:**
- If NO course content is provided in context, do NOT make up topics
- Only explain topics that are explicitly mentioned in the course materials
- Don't assume the user is asking about a specific subject unless clearly indicated
- Greetings like "kesa hai" should get friendly responses, NOT academic lectures"""

def chat_with_context(user_message, session_id, user_id):
    """Main chat function with conversation memory and intelligent context"""
    
    # Get or create session
    chat_session = get_or_create_session(session_id, user_id)
    if chat_session is None:
        return None
    history = get_conversation_history(session_id, user_id, last_n=6)
    
    # ✅ FIXED: Detect intent FIRST
    intent = detect_intent(user_message)
    
    # ✅ FIXED: Skip RAG for non-academic intents
    if intent in ["greeting", "farewell", "acknowledgment"]:
        assistant_reply = handle_casual_conversation(user_message, intent)
        
        # Still save to session for continuity
        update_session_context(session_id, user_id, user_message, assistant_reply, None)
        chats_collection.insert_one({
            "session_id": session_id,
            "user_id": user_id,
            "user_message": user_message,
            "assistant_reply": assistant_reply,
            "detected_subject": None,
            "intent": intent,
            "timestamp": datetime.now()
        })
        
        print(f"[INTENT: {intent}] No RAG needed")
        return assistant_reply
    
    # Only proceed with RAG for academic questions
    recent_context = " ".join([msg["content"] for msg in history[-3:]])
    full_context = f"{recent_context} {user_message}"
    detected_subject = detect_subject_from_text(full_context)
    
    if not detected_subject and chat_session.get("context", {}).get("current_subject"):
        detected_subject = chat_session["context"]["current_subject"]
    
    # Build conversation history for LLM
    conversation_str = ""
    if history:
        conversation_str = "**Previous conversation:**\n"
        for msg in history[-4:]:
            role = "Student" if msg["role"] == "user" else "Learnify"
            conversation_str += f"{role}: {msg['content']}\n"
        conversation_str += "\n"
    
    # FIXED: Only query RAG for academic questions
    rag_context = ""
    document_context = ""
    
    # Only RAG if subject detected OR question intent
    if detected_subject or intent == "question":
        where_filter = {"course_code": detected_subject} if detected_subject else None
        
        try:
            results = collection.query(
                query_texts=[user_message],
                n_results=5,
                where=where_filter
            )
            
            if results['documents'] and results['documents'][0]:
                rag_context = "**Relevant course topics:**\n"
                for i, doc in enumerate(results['documents'][0]):
                    metadata = results['metadatas'][0][i]
                    rag_context += f"\n{metadata['course']} - Week {metadata['week']}:\n{doc}\n"
        except Exception as e:
            print(f"ChromaDB query error: {e}")

    document_context = build_document_context(session_id, user_message)
    
    # Build enhanced prompt
    enhanced_context = SYSTEM_PROMPT
    
    if conversation_str:
        enhanced_context += f"\n\n{conversation_str}"
    
    if detected_subject:
        enhanced_context += f"\n**Current subject focus:** {detected_subject}\n"
    
    # ✅ FIXED: Only add RAG context if it exists
    if rag_context:
        enhanced_context += f"\n{rag_context}\n"
    else:
        enhanced_context += "\n**Note:** No specific course content found for this query. Answer based on general knowledge if appropriate.\n"

    if document_context:
        enhanced_context += f"\n{document_context}\n"
        enhanced_context += (
            "\n**Priority instruction:** If the uploaded document context answers the question, "
            "prefer that context over general assumptions.\n"
        )
    
    enhanced_context += f"\n**Student's current question:** {user_message}\n\n"
    enhanced_context += (
        "**Instructions:** Answer in Roman Urdu + English mix. Be specific and helpful. "
        "Reference the course topics if relevant. For academic answers, always return clean "
        "Markdown-style formatting with headings, bullets, spacing, and bold key terms."
    )
    
    # Debug logging
    print(f"[INTENT: {intent}] [SUBJECT: {detected_subject}] [RAG: {bool(rag_context)}]")
    
    # Call Gemini LLM
    try:
        assistant_reply = generate_with_gemini(enhanced_context, user_message)
        # Emergency Groq fallback:
        # assistant_reply = generate_with_groq(enhanced_context, user_message)
        
        # Update session with new context
        update_session_context(session_id, user_id, user_message, assistant_reply, detected_subject)
        
        # Also save to chats collection for history
        chats_collection.insert_one({
            "session_id": session_id,
            "user_id": user_id,
            "user_message": user_message,
            "assistant_reply": assistant_reply,
            "detected_subject": detected_subject,
            "intent": intent,
            "timestamp": datetime.now()
        })
        
        return assistant_reply
        
    except Exception as e:
        return f"Error: {str(e)}"

# ========== FLASK ROUTES ==========
@app.route('/chat', methods=['POST'])
def chat():
    """Main chat endpoint"""
    user_id, auth_error = require_current_user()
    if auth_error:
        return auth_error

    data = request.json
    user_message = data.get('question', '')
    session_id = data.get('session_id', str(uuid.uuid4()))
    
    if not user_message:
        return jsonify({"error": "No question provided"}), 400
    
    response = chat_with_context(user_message, session_id, user_id)
    if response is None:
        return jsonify({"error": "Session not found"}), 404
    
    return jsonify({
        "answer": response,
        "session_id": session_id
    })

@app.route('/new-chat', methods=['POST'])
def new_chat():
    """Return a new chat id without saving an empty conversation."""
    user_id, auth_error = require_current_user()
    if auth_error:
        return auth_error

    session_id = str(uuid.uuid4())
    return jsonify({"session_id": session_id})


@app.route('/upload-document', methods=['POST'])
def upload_document():
    """Upload a session document and index it for chat retrieval."""
    user_id, auth_error = require_current_user()
    if auth_error:
        return auth_error

    session_id = request.form.get('session_id', '').strip()
    upload = request.files.get('file')

    if not session_id:
        return jsonify({"error": "Session ID is required"}), 400

    if not sessions_collection.find_one(session_owner_filter(session_id, user_id)):
        return jsonify({"error": "Session not found"}), 404

    if upload is None or not upload.filename:
        return jsonify({"error": "No file uploaded"}), 400

    extension = Path(upload.filename).suffix.lower()
    if extension not in ALLOWED_DOCUMENT_EXTENSIONS:
        return jsonify({"error": "Only PDF, DOCX, and TXT files are supported"}), 400

    safe_name = secure_filename(upload.filename)
    document_id = str(uuid.uuid4())
    stored_name = f"{session_id}_{document_id}_{safe_name}"
    saved_path = UPLOADS_FOLDER / stored_name

    try:
        upload.save(saved_path)
        extracted_text = extract_document_text(saved_path)
        if not extracted_text:
            raise RuntimeError("No readable text found in uploaded document.")

        chunks = chunk_text(extracted_text)
        if not chunks:
            raise RuntimeError("Uploaded document did not produce usable text chunks.")

        ids = []
        documents = []
        metadatas = []
        for index, chunk in enumerate(chunks, start=1):
            ids.append(f"{document_id}_chunk_{index}")
            documents.append(chunk)
            metadatas.append({
                "session_id": session_id,
                "document_id": document_id,
                "file_name": safe_name,
                "chunk_index": index,
            })

        session_docs_collection.add(ids=ids, documents=documents, metadatas=metadatas)

        document_info = {
            "document_id": document_id,
            "file_name": safe_name,
            "extension": extension,
            "chunk_count": len(chunks),
            "uploaded_at": datetime.now().isoformat(),
        }
        add_document_to_session(session_id, user_id, document_info)

        return jsonify({
            "success": True,
            "document": document_info,
            "documents": get_session_documents(session_id, user_id),
        })
    except Exception as exc:
        if saved_path.exists():
            saved_path.unlink(missing_ok=True)
        remove_session_document_embeddings(session_id, document_id)
        return jsonify({"error": str(exc)}), 500


@app.route('/documents/<session_id>', methods=['GET'])
def get_documents(session_id):
    """Return uploaded documents for a chat session."""
    user_id, auth_error = require_current_user()
    if auth_error:
        return auth_error

    chat_session = sessions_collection.find_one(
        session_owner_filter(session_id, user_id),
        {"documents": 1, "_id": 0},
    )
    if not chat_session:
        return jsonify({"error": "Session not found"}), 404
    return jsonify({"documents": chat_session.get("documents", [])})


@app.route('/documents/<session_id>/<document_id>', methods=['DELETE'])
def delete_document(session_id, document_id):
    """Remove an uploaded document from a chat session."""
    user_id, auth_error = require_current_user()
    if auth_error:
        return auth_error

    chat_session = sessions_collection.find_one(session_owner_filter(session_id, user_id))
    if not chat_session:
        return jsonify({"error": "Session not found"}), 404

    documents = chat_session.get("documents", [])
    target = next((doc for doc in documents if doc.get("document_id") == document_id), None)
    if not target:
        return jsonify({"error": "Document not found"}), 404

    remove_session_document_embeddings(session_id, document_id)

    file_prefix = f"{session_id}_{document_id}_"
    for path in UPLOADS_FOLDER.glob(f"{file_prefix}*"):
        path.unlink(missing_ok=True)

    sessions_collection.update_one(
        session_owner_filter(session_id, user_id),
        {
            "$pull": {"documents": {"document_id": document_id}},
            "$set": {"updated_at": datetime.now()},
        },
    )

    return jsonify({"success": True, "documents": get_session_documents(session_id, user_id)})

@app.route('/chat-history/<session_id>', methods=['GET'])
def get_chat_history(session_id):
    """Get full chat history for a session"""
    user_id, auth_error = require_current_user()
    if auth_error:
        return auth_error

    chat_session = sessions_collection.find_one(session_owner_filter(session_id, user_id))
    if not chat_session:
        return jsonify({"error": "Session not found"}), 404
    
    messages = chat_session.get("messages", [])
    for msg in messages:
        msg["timestamp"] = msg["timestamp"].isoformat()
    
    return jsonify({"messages": messages})

@app.route('/sessions', methods=['GET'])
def get_sessions():
    """Get chat sessions for the logged-in user."""
    user_id, auth_error = require_current_user()
    if auth_error:
        return auth_error

    sessions = list(
        sessions_collection.find(
            {"user_id": user_id, "messages.0": {"$exists": True}},
            {"_id": 0},
        ).sort("created_at", -1)
    )
    
    for session in sessions:
        session["created_at"] = session["created_at"].isoformat()
        if "updated_at" in session:
            session["updated_at"] = session["updated_at"].isoformat()
    
    return jsonify({"sessions": sessions})

@app.route('/delete-chat/<session_id>', methods=['DELETE'])
def delete_chat(session_id):
    """Delete a chat session"""
    user_id, auth_error = require_current_user()
    if auth_error:
        return auth_error

    result = sessions_collection.delete_one(session_owner_filter(session_id, user_id))
    if result.deleted_count > 0:
        return jsonify({"message": "Chat deleted successfully"})
    return jsonify({"error": "Session not found"}), 404

@app.route('/health', methods=['GET'])
def health():
    """Health check"""
    user_id = get_current_user_id()
    active_sessions_filter = (
        {"user_id": user_id, "messages.0": {"$exists": True}}
        if user_id
        else {"user_id": "__none__"}
    )
    return jsonify({
        "status": "ok",
        "chroma_count": collection.count(),
        "active_sessions": sessions_collection.count_documents(active_sessions_filter)
    })

# ========== MAIN ==========
if __name__ == "__main__":
    print("="*80)
    print("LEARNIFY INTELLIGENT BACKEND - PRODUCTION READY")
    print("="*80)
    print("\nStep 1: Processing course files...")
    courses = process_course_files()
    
    if courses:
        print("\n" + "="*80)
        print("Step 2: Starting Flask Server")
        print("="*80)
        print("\n✓ Server: http://localhost:5000")
        print("✓ MongoDB: Connected")
        print("✓ ChromaDB: Loaded")
        print("✓ Gemini LLM: Ready")
        print("✓ Intent Detection: Active")
        print("\nLearnify is ready! 🚀\n")
        
        app.run(debug=False, use_reloader=False, host='0.0.0.0', port=5000)
    else:
        print("\n❌ Failed to load courses. Check 'files' folder.")
